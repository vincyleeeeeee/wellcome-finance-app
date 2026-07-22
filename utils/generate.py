"""Document generation: confirmation letter (docx) and invoice (xlsx)."""

import os
import datetime
from typing import Tuple
from docx import Document
from docx.oxml.ns import qn
import openpyxl

# Template paths (relative to Wellcome project)
TEMPLATE_DIR = "/Users/vincy/Documents/Wellcome/项目/_模板"
OUTPUT_BASE = "/Users/vincy/Documents/Wellcome/项目"

# Chinese number mapping
_CN_DIGITS = ["零", "壹", "贰", "叁", "肆", "伍", "陆", "柒", "捌", "玖"]
_CN_UNITS = ["", "拾", "佰", "仟", "万"]
_CN_UNITS_BIG = ["", "万", "亿"]


def _num_to_chinese(n: int) -> str:
    """Convert integer to Chinese uppercase number. E.g., 5400 -> 伍仟肆佰"""
    if n == 0:
        return "零"
    s = str(n)
    length = len(s)
    result = ""
    prev_zero = False
    for i, ch in enumerate(s):
        digit = int(ch)
        pos = length - i - 1
        unit_idx = pos % 4
        if digit == 0:
            prev_zero = True
            if unit_idx == 0:  # at 万/亿 boundary
                big_idx = pos // 4
                if big_idx > 0:
                    result += _CN_UNITS_BIG[big_idx]
                prev_zero = False
        else:
            if prev_zero:
                result += "零"
                prev_zero = False
            result += _CN_DIGITS[digit] + _CN_UNITS[unit_idx]
            if unit_idx == 0:
                big_idx = pos // 4
                if big_idx > 0:
                    result += _CN_UNITS_BIG[big_idx]
    return result


def _amount_chinese(amount: int, currency: str = "USD") -> str:
    """Get Chinese uppercase amount string with currency suffix."""
    cn = _num_to_chinese(amount)
    if currency == "RMB":
        return f"{cn}元整"
    return f"{cn}元整"


def generate_confirmation_letter(client: dict, project: dict) -> str:
    """
    Generate confirmation letter .docx.
    Returns the file path of the generated document.
    """
    output_dir = os.path.join(OUTPUT_BASE, project['client_short'], project['brand_name'], "财务")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{project['brand_name']}-{project['total_posts'].replace(' ', '-')}-Confirmation-Letter.docx"
    output_path = os.path.join(output_dir, filename)

    doc = Document(os.path.join(TEMPLATE_DIR, "Confirmation-Letter-Template.docx"))

    currency_symbol = "RMB" if project.get('currency') == "RMB" else "USD"
    amount_str = f"{currency_symbol} {project['amount']:,.2f}"

    # Table 0: submitter info
    t0 = doc.tables[0]
    cell00 = t0.rows[0].cells[0]
    cell00.paragraphs[1].runs[0].text = f"   {client['contact']}    "
    cell00.paragraphs[3].runs[0].text = client['full_name']
    cell00.paragraphs[4].runs[0].text = f"申請日期/Date：  {project['application_date']}         "

    # Table 1: project details
    t1 = doc.tables[1]
    t1.rows[0].cells[1].paragraphs[0].runs[0].text = project['project_code']
    t1.rows[1].cells[1].paragraphs[0].runs[0].text = project['project_name']
    t1.rows[2].cells[1].paragraphs[0].runs[0].text = amount_str
    t1.rows[3].cells[1].paragraphs[0].runs[0].text = project['brand_name']
    t1.rows[4].cells[1].paragraphs[0].runs[0].text = project['venue']
    t1.rows[5].cells[1].paragraphs[0].runs[0].text = project['execution_period']
    t1.rows[6].cells[1].paragraphs[0].runs[0].text = project['shooting_date']
    t1.rows[7].cells[1].paragraphs[0].runs[0].text = project['total_posts']

    # Unify fonts
    TARGET_FONT = 'PingFang SC'
    for p in doc.paragraphs:
        for r in p.runs:
            r.font.name = TARGET_FONT
            r._element.rPr.rFonts.set(qn('w:eastAsia'), TARGET_FONT)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = TARGET_FONT
                        r._element.rPr.rFonts.set(qn('w:eastAsia'), TARGET_FONT)

    doc.save(output_path)
    return output_path


def generate_invoice(client: dict, project: dict) -> str:
    """
    Generate invoice .xlsx.
    Returns the file path of the generated document.
    """
    output_dir = os.path.join(OUTPUT_BASE, project['client_short'], project['brand_name'], "财务")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{project['brand_name']}-{project['total_posts'].replace(' ', '-')}-invoice.xlsx"
    output_path = os.path.join(output_dir, filename)

    wb = openpyxl.load_workbook(os.path.join(TEMPLATE_DIR, "Invoice-Template.xlsx"))
    ws = wb.active

    currency = project.get('currency', 'USD')
    is_rmb = currency == "RMB"
    currency_label = "RMB" if is_rmb else "USD"
    amount = project['amount']

    # Header
    ws['C3'] = project.get('invoice_project_name', project['project_name'])
    ws['C4'] = project['execution_period']
    ws['C5'] = project['venue']
    ws['C7'] = client['full_name']
    ws['C8'] = client['address']
    ws['C9'] = client['contact']
    ws['C10'] = client.get('phone') if client.get('phone') and client['phone'] != '（待补充）' else None
    ws['C11'] = client.get('email') if client.get('email') and client['email'] != '（待补充）' else None
    ws['E8'] = project['project_code']
    ws['E9'] = project['invoice_date']
    ws['E10'] = project['due_date']
    ws['E11'] = project['project_code']

    # Currency headers
    if is_rmb:
        ws['D13'] = '單價（人民幣）\nUNIT PRICE (RMB)'
        ws['G13'] = '小計（人民幣）\nAMOUNT (RMB)'
    else:
        ws['D13'] = '單價（美元）\nUNIT PRICE (USD)'
        ws['G13'] = '小計（美元）\nAMOUNT (USD)'

    # Line items
    ws['D15'] = amount
    ws['E15'] = 1
    ws['G15'] = amount

    # Description
    cn_amount = _amount_chinese(int(amount), currency)
    ws['C16'] = "項目「   服务      」款\nltem \"Service'"
    ws['C18'] = f"總付款金額為{cn_amount}\nFull payment of {currency_label}, {amount:,}"

    wb.save(output_path)
    return output_path


def generate_email_confirmation(project: dict) -> Tuple[str, str, str]:
    """
    Generate email copy for confirmation letter.
    Returns (subject, body, recipient_hint).
    """
    currency = project.get('currency', 'USD')
    currency_symbol = "¥" if currency == "RMB" else "$"
    amount_display = f"{currency_symbol}{project['amount']:,}"

    subject = f"请确认：{project['brand_name']} {project['project_name']} 合作细节 / {project['project_code']}"

    body = f"""感谢您确认与我司就 **【{project['brand_name']} {project['project_name']}】** 开展合作。为推进后续执行，现将双方沟通确认的活动细节汇总如下：

**【合作细节】**
- 合作内容：{project.get('content_type', 'UGC铺量')}
- 内容数量：{project['total_posts']}
- 合作费用：{amount_display}
- 发布平台：{project.get('platform', '小红书')}
- 执行周期：{project['execution_period']}

**【付款安排】**
以上细节无误后，我们将开具 INVOICE。

请您回复本邮件确认以下两点：
1. 上述活动细节无误；
2. 贵方可预期的付款时间（或时间范围）。"""

    return subject, body


def generate_email_invoice(project: dict) -> Tuple[str, str, str]:
    """
    Generate email copy for invoice delivery.
    Returns (subject, body, recipient_hint).
    """
    currency = project.get('currency', 'USD')
    currency_symbol = "¥" if currency == "RMB" else "$"
    amount_display = f"{currency_symbol}{project['amount']:,}"

    subject = f"Invoice 请查收 — {project['brand_name']} {project['project_name']} / {project['project_code']}"

    body = f"""附件为本项目 Invoice，请查收。

| 项目 | 内容 |
|------|------|
| Invoice No. | {project['project_code']} |
| 项目 | {project['project_name']} |
| 金额 | {currency} {amount_display} |
| 到期日 | {project['due_date']} |

请于到期日前安排付款，如有疑问请随时联系。谢谢！"""

    return subject, body


# ============================================================
# Cash Receipt generation
# ============================================================

# Fixed Wellcome issuer info
ISSUER = {
    "name": "Mr. Terry.Su",
    "phone": "008613609023860",
    "address": "UNIT 1021, BEVERLEY COMMERCIAL CENTRE, 87-105 CHATHAN ROAD SOUTH, TSIM SHA TSUI, HK",
    "company": "WELLCOME (INTERNATIONAL) LIMITED",
}


def generate_cash_receipt(client: dict, receipt_data: dict) -> str:
    """
    Generate cash receipt .xlsx from template.
    Returns the file path of the generated receipt.
    """
    output_dir = os.path.join(OUTPUT_BASE, receipt_data.get('client_short', ''),
                              receipt_data.get('brand_name', ''), "财务")
    os.makedirs(output_dir, exist_ok=True)

    filename = f"{receipt_data['brand_name']}-Receipt-{receipt_data['project_code']}.xlsx"
    output_path = os.path.join(output_dir, filename)

    wb = openpyxl.load_workbook(os.path.join(TEMPLATE_DIR, "Cash-Receipt-Template.xlsx"))
    ws = wb.active

    currency = receipt_data.get('currency', 'USD')
    currency_label = "RMB" if currency == "RMB" else "USD"
    amount = receipt_data['amount']
    payment_amount = receipt_data.get('payment_amount', amount)

    # Header - left side
    ws['C3'] = receipt_data.get('project_name', '')
    ws['C4'] = receipt_data.get('project_date', '')
    ws['C5'] = receipt_data.get('venue', '')
    ws['C7'] = client['full_name']
    ws['C8'] = client.get('address', '')
    ws['C9'] = client['contact']
    ws['C10'] = client.get('phone') if client.get('phone') and client['phone'] != '（待补充）' else ''
    ws['C11'] = client.get('email') if client.get('email') and client['email'] != '（待补充）' else ''

    # Header - right side
    ws['E3'] = receipt_data.get('issuer_name', ISSUER['name'])
    ws['E4'] = ISSUER['phone']
    ws['E5'] = ISSUER['address']
    ws['E8'] = receipt_data.get('project_code', '')
    ws['E9'] = receipt_data.get('payment_date', datetime.datetime.now())
    ws['E10'] = receipt_data.get('gained_date', datetime.datetime.now())
    ws['E11'] = receipt_data.get('payment_method', 'BANK')

    # Body text
    ws['C13'] = (f"Received From {ISSUER['company']} The amount of "
                 f"{currency_label} {payment_amount:,.2f}\n"
                 f"For the {receipt_data.get('project_name', '')} Project")

    # Signature area
    gained_date = receipt_data.get('gained_date', datetime.datetime.now())
    if isinstance(gained_date, datetime.datetime):
        date_str = gained_date.strftime('%Y/%m/%d')
    else:
        date_str = str(gained_date)
    ws['D15'] = f"Name：\n\nDate：{date_str}\n\nSignature：\n"

    wb.save(output_path)
    return output_path
